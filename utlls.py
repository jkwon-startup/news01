import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import base64

@st.cache_data
def search_naver_news(keyword, num_articles):
    url = f"https://search.naver.com/search.naver?where=news&query={keyword}&sm=tab_jum&sort=0"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'lxml')

    news_list = []
    seen_titles = set()

    news_elements = soup.select("ul.list_news > li")
    for element in news_elements[:num_articles]:
        title = element.select_one("a.news_tit").get_text()
        summary = element.select_one("div.news_dsc").get_text()
        link = element.select_one("a.news_tit")['href']

        if title not in seen_titles:
            seen_titles.add(title)
            news_list.append({
                'title': title,
                'summary': summary,
                'link': link
            })

    return news_list

def create_word_document(news_list, keyword):
    doc = Document()
    doc.add_heading(f"'{keyword}' 관련 뉴스", 0)

    for idx, news in enumerate(news_list, 1):
        doc.add_heading(f"{idx}. {news['title']}", level=1)
        doc.add_paragraph(news['summary'])
        doc.add_paragraph(f"링크: {news['link']}")
        doc.add_paragraph("\n")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_download_link(buffer, filename):
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">다운로드</a>'